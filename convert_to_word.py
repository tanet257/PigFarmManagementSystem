#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_table_from_markdown(doc, markdown_table):
    """Convert markdown table to Word table"""
    lines = [line.strip() for line in markdown_table.strip().split('\n') if line.strip()]

    # Skip header separator
    rows_data = []
    for i, line in enumerate(lines):
        if i == 1:  # Skip the separator line
            continue
        if line.startswith('|') and line.endswith('|'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            rows_data.append(cells)

    if not rows_data:
        return

    # Create table
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.style = 'Light Grid Accent 1'

    # Add data to table
    for i, row_data in enumerate(rows_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            # Format header row
            if i == 0:
                for paragraph in row.cells[j].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)

def read_markdown_file(filepath):
    """Read markdown file and extract content"""
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()

def convert_md_to_word(md_content):
    """Convert markdown content to Word document"""
    doc = Document()

    # Add title
    title = doc.add_heading('ระบบจัดการฟาร์มหมู', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_heading('03 API Design, Backend Development, and Testing', level=1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Split content into lines
    lines = md_content.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading 1
        if line.startswith('## '):
            heading_text = line.replace('## ', '').strip()
            doc.add_heading(heading_text, level=1)
            i += 1
        # Heading 2
        elif line.startswith('### '):
            heading_text = line.replace('### ', '').strip()
            doc.add_heading(heading_text, level=2)
            i += 1
        # Heading 3
        elif line.startswith('#### '):
            heading_text = line.replace('#### ', '').strip()
            doc.add_heading(heading_text, level=3)
            i += 1
        # Code block
        elif line.strip().startswith('```'):
            # Collect code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                code_para = doc.add_paragraph()
                code_para.style = 'Normal'
                code_run = code_para.add_run('\n'.join(code_lines))
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
            i += 1
        # Table
        elif line.strip().startswith('|'):
            # Collect full table
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if table_lines:
                add_table_from_markdown(doc, '\n'.join(table_lines))
        # Bullet point
        elif line.strip().startswith('-') and not line.strip().startswith('---'):
            bullet_text = line.strip()[1:].strip()
            para = doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1
        # Numbered list
        elif re.match(r'^\s*\d+\.', line):
            match = re.match(r'^\s*\d+\.\s*(.*)', line)
            if match:
                list_text = match.group(1)
                doc.add_paragraph(list_text, style='List Number')
            i += 1
        # Empty line
        elif not line.strip():
            doc.add_paragraph()
            i += 1
        # Regular paragraph
        elif line.strip() and not line.strip().startswith('#'):
            para = doc.add_paragraph(line)
            i += 1
        else:
            i += 1

    return doc

# Main execution
md_file = 'c:\\Users\\thete\\Documents\\GitHub\\PigFarmManagementSystem\\DOCUMENTATION\\03_API_BACKEND_TESTING.md'
word_file = 'c:\\Users\\thete\\Documents\\GitHub\\PigFarmManagementSystem\\DOCUMENTATION\\03_API_BACKEND_TESTING.docx'

print(f"Reading markdown file: {md_file}")
md_content = read_markdown_file(md_file)

print(f"Converting to Word document...")
doc = convert_md_to_word(md_content)

print(f"Saving to: {word_file}")
doc.save(word_file)

print("Done! Word document created successfully.")
print(f"File saved at: {word_file}")
